from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "Privacy_Control_Center_July_9_2026_Work_Summary.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color="D9DEE7"):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_width(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            cell = row.cells[idx]
            cell.width = Inches(width)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text)
    return p


def add_status_table(doc, rows):
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    headers = ["Area", "What Changed", "Status"]
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, "E8EEF5")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(header)
        run.bold = True
    for area, change, status in rows:
        cells = table.add_row().cells
        cells[0].text = area
        cells[1].text = change
        cells[2].text = status
    set_table_width(table, [1.55, 3.85, 1.10])
    set_table_borders(table)
    doc.add_paragraph()


def configure_styles(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    title = styles["Title"]
    title.font.name = "Calibri"
    title.font.size = Pt(20)
    title.font.bold = True
    title.font.color.rgb = RGBColor(31, 58, 95)
    title.paragraph_format.space_after = Pt(6)

    for name, size, color, before, after in (
        ("Heading 1", 16, RGBColor(46, 116, 181), 16, 8),
        ("Heading 2", 13, RGBColor(46, 116, 181), 12, 6),
        ("Heading 3", 12, RGBColor(31, 77, 120), 8, 4),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def main():
    doc = Document()
    configure_styles(doc)

    title = doc.add_paragraph(style="Title")
    title.add_run("Privacy Control Center - July 9, 2026 Work Summary")
    subtitle = doc.add_paragraph()
    subtitle.add_run("Project: React + FastAPI + Databricks REST API").bold = True
    subtitle.add_run(" | Backup updated: privacy-control-center-20260709-111545.zip")

    doc.add_heading("Executive Summary", level=1)
    doc.add_paragraph(
        "Today focused on validating and strengthening Unity Catalog permission workflows, "
        "adding group-management capabilities, improving frontend permission search behavior, "
        "and fixing the table Excel export path."
    )

    add_status_table(doc, [
        (
            "Permission Writes",
            "Verified Unity Catalog permission read/write behavior through Databricks REST API and integrated grant, update, and remove flows.",
            "Completed",
        ),
        (
            "Privilege Loading",
            "Replaced unavailable Databricks dynamic privilege discovery with a centralized backend privilege registry.",
            "Completed",
        ),
        (
            "Grant Dialog",
            "Kept the existing dialog design while adding user and group-aware titles, principal labels, validation, and audit events.",
            "Completed",
        ),
        (
            "Group Management",
            "Added a SCIM-backed Group Management module for listing groups, viewing members, creating/editing/deleting groups, and membership changes.",
            "Completed",
        ),
        (
            "Table Export",
            "Fixed Excel export when table statistics are unavailable by safely handling null statistics.",
            "Completed",
        ),
        (
            "Workspace Groups vs UC Principals",
            "Identified that workspace SCIM groups can exist without being accepted as Unity Catalog grant principals.",
            "Known limitation",
        ),
    ])

    doc.add_heading("Backend Changes", level=1)
    backend_items = [
        "Added catalog permission write route: PATCH /catalogs/{catalog_name}/permissions.",
        "Added reusable Unity Catalog permission helpers for grant, update, and remove.",
        "Added centralized available-privilege registry for catalog, schema, and table securables.",
        "Added SCIM group helpers for listing, reading, creating, updating, deleting, adding members, and removing members.",
        "Added clearer error handling for unsupported APIs, missing principals, permission denial, and Free Edition limitations.",
        "Added group principal resolution logic to distinguish Unity Catalog grant principals from workspace-only SCIM groups.",
    ]
    for item in backend_items:
        add_bullet(doc, item)

    doc.add_heading("Frontend Changes", level=1)
    frontend_items = [
        "Added PermissionDialog reuse for user and group grants without duplicating privilege-loading logic.",
        "Extended Search Group behavior to mirror Search User for existing Unity Catalog permission rows.",
        "Added group-specific empty states and controlled Grant Group Access behavior.",
        "Added Group Management page, route, navigation entry, modals, audit counter, and Excel export support.",
        "Improved user/group suggestions behavior and logging for failed API calls.",
        "Fixed table workbook export when statistics are missing or still loading.",
    ]
    for item in frontend_items:
        add_bullet(doc, item)

    doc.add_heading("Databricks Findings", level=1)
    findings = [
        "The current workspace can read Unity Catalog permissions for the tested catalog.",
        "A prior write test confirmed PATCH permission updates can succeed for supported principals and privileges.",
        "Databricks does not expose a REST endpoint that returns the full list of available Unity Catalog privileges for a securable.",
        "Workspace SCIM groups such as 'new' may be manageable through SCIM but still not grantable in Unity Catalog permissions.",
        "The built-in Unity Catalog group principal 'account users' appears in catalog permissions and remains the reliable tested group principal.",
    ]
    for item in findings:
        add_bullet(doc, item)

    doc.add_heading("Validation Performed", level=1)
    validations = [
        "Ran Python compile checks for backend app and Databricks service modules.",
        "Ran Vite production builds after frontend changes.",
        "Restarted FastAPI where stale local servers were serving older routes.",
        "Verified read-only API routes for catalog permissions, groups, group details, and available privileges.",
        "Avoided destructive Databricks tests for group create/delete and permission changes unless already explicitly requested.",
    ]
    for item in validations:
        add_bullet(doc, item)

    doc.add_heading("Files and Backup", level=1)
    doc.add_paragraph("Updated backup:")
    add_bullet(doc, r"C:\Users\harih\Privacy-Control-Center\backups\privacy-control-center-20260709-111545")
    add_bullet(doc, r"C:\Users\harih\Privacy-Control-Center\backups\privacy-control-center-20260709-111545.zip")
    doc.add_paragraph("Primary source areas included in the backup:")
    for item in [
        "backend/app.py",
        "backend/services/databricks_api.py",
        "frontend/src",
        "frontend/package.json and package-lock.json",
        "frontend/vite.config.js and index.html",
    ]:
        add_bullet(doc, item)

    doc.add_heading("Follow-Up Notes", level=1)
    follow_up = [
        "Resolve Unity Catalog account-level group creation or assignment if custom groups need to be grantable.",
        "Optionally add a backend endpoint that explicitly lists only grantable Unity Catalog principals.",
        "Consider chunk splitting in the frontend build later; Vite reports a chunk-size warning but the build succeeds.",
        "Re-test real group permission grant/update/remove once an account-level grantable group is available.",
    ]
    for item in follow_up:
        add_bullet(doc, item)

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run(f"Generated {datetime.now().strftime('%d %b %Y, %H:%M')}")

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
