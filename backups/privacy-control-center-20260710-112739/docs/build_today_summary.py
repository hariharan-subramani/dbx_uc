from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path(__file__).with_name("Privacy_Control_Center_Today_Summary.docx")
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "202A31"
MUTED = "66727B"
LIGHT = "F2F4F7"
RED = "D53F3F"
GREEN = "2B7A57"
WHITE = "FFFFFF"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths[index]
            cell.width = Inches(width / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_font(run, size=11, color=INK, bold=False, italic=False):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic


def add_field(paragraph, instruction):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.1
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    set_font(p.add_run(text))
    return p


def add_step(doc, title, detail):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(7)
    p.paragraph_format.line_spacing = 1.1
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    set_font(p.add_run(title + " "), bold=True)
    set_font(p.add_run(detail))


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        set_cell_shading(cell, LIGHT)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        set_font(p.add_run(header), size=10, color=DARK_BLUE, bold=True)
    for row_data in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row_data):
            p = cells[index].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05
            set_font(p.add_run(str(value)), size=9.5)
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(0.8)
section.bottom_margin = Inches(0.75)
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.header_distance = Inches(0.492)
section.footer_distance = Inches(0.492)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
normal.font.size = Pt(11)
normal.font.color.rgb = RGBColor.from_string(INK)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.1

for name, size, color, before, after in (
    ("Heading 1", 16, BLUE, 16, 8),
    ("Heading 2", 13, BLUE, 12, 6),
    ("Heading 3", 12, DARK_BLUE, 8, 4),
):
    style = styles[name]
    style.font.name = "Calibri"
    style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = RGBColor.from_string(color)
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.keep_with_next = True

header = section.header
hp = header.paragraphs[0]
hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
set_font(hp.add_run("PRIVACY CONTROL CENTER  /  DEVELOPMENT SUMMARY"), size=8.5, color=MUTED, bold=True)

footer = section.footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
set_font(fp.add_run("June 30, 2026  |  Page "), size=8.5, color=MUTED)
add_field(fp, "PAGE")

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(12)
p.paragraph_format.space_after = Pt(3)
set_font(p.add_run("IMPLEMENTATION SUMMARY"), size=9, color=RED, bold=True)
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(4)
set_font(p.add_run("Privacy & Compliance Control Center"), size=25, color=INK, bold=True)
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(15)
set_font(p.add_run("Databricks-inspired frontend, Unity Catalog drill-down, and live data preview"), size=13, color=MUTED)

metadata = [
    ("Project", "Privacy-Control-Center"),
    ("Date", "June 30, 2026"),
    ("Stack", "React 19 + Vite 8 / FastAPI / Databricks SDK"),
    ("Outcome", "Working catalog → schema → table → row-data exploration"),
]
for label, value in metadata:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    set_font(p.add_run(f"{label}: "), size=10.5, bold=True)
    set_font(p.add_run(value), size=10.5)

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(12)
p.paragraph_format.space_after = Pt(0)
p_pr = p._p.get_or_add_pPr()
border = OxmlElement("w:pBdr")
bottom = OxmlElement("w:bottom")
bottom.set(qn("w:val"), "single")
bottom.set(qn("w:sz"), "12")
bottom.set(qn("w:space"), "1")
bottom.set(qn("w:color"), RED)
border.append(bottom)
p_pr.append(border)

doc.add_heading("Executive summary", level=1)
p = doc.add_paragraph()
set_font(p.add_run(
    "The project evolved from a simple centered catalog card into a professional application shell modeled on the density and interaction patterns of Databricks. "
    "The completed experience connects to the existing Databricks workspace and supports progressive, in-place exploration from catalogs to schemas, tables, and a read-only data preview."
))

add_table(
    doc,
    ["Area", "Result", "Status"],
    [
        ("Frontend", "Professional responsive dashboard and nested accordion browser", "Complete"),
        ("Metadata APIs", "Catalog, schema, and table discovery through Databricks SDK", "Verified"),
        ("Data access", "SQL warehouse-backed preview capped at 50 rows", "Verified"),
        ("Runtime", "Frontend and backend active on local development ports", "Running"),
    ],
    [1800, 5940, 1620],
)

doc.add_heading("1. Frontend modernization", level=1)
p = doc.add_paragraph()
set_font(p.add_run(
    "The original page used a single card with inline styles and residual Vite starter CSS. It was replaced with a cohesive application layout and reusable class-based styling."
))
for item in (
    "Added a fixed workspace sidebar, top search bar, user profile area, breadcrumbs, page heading, metric cards, and catalog panel.",
    "Adopted a restrained red accent, neutral surfaces, compact typography, SVG interface icons, and Databricks-like information density.",
    "Added responsive behavior for tablet and mobile widths, including a collapsed navigation rail and simplified table layouts.",
    "Introduced consistent loading skeletons, empty states, API errors, retry actions, selected states, and accessible expansion attributes.",
):
    add_bullet(doc, item)

doc.add_heading("2. Progressive data exploration", level=1)
p = doc.add_paragraph()
set_font(p.add_run(
    "The main catalog panel now works as a four-level nested explorer. Every level expands directly beneath the selected item, avoiding separate detail pages and preserving context."
))
add_table(
    doc,
    ["Level", "User action", "Displayed content"],
    [
        ("Catalog", "Click a catalog row", "Schemas visible to the workspace identity"),
        ("Schema", "Click a schema row", "Tables and views within that schema"),
        ("Table", "Click a table row", "First 25 records in a scrollable grid"),
        ("Any open level", "Click the selected row again", "The level collapses"),
    ],
    [1500, 2700, 5160],
)

doc.add_heading("Data preview behavior", level=2)
for item in (
    "Shows column names and Databricks data types in sticky grid headers.",
    "Displays up to 25 rows in the UI; the backend enforces an absolute maximum of 50.",
    "Formats missing values as NULL and provides horizontal and vertical scrolling for wide datasets.",
    "Uses a read-only SELECT statement with safely quoted catalog, schema, and table identifiers.",
    "Automatically selects a running SQL warehouse, or falls back to the first available warehouse.",
):
    add_bullet(doc, item)

doc.add_heading("3. Backend API additions", level=1)
p = doc.add_paragraph()
set_font(p.add_run(
    "FastAPI routes were added around the Databricks SDK while preserving the existing user and catalog endpoints. CORS now accepts both localhost and 127.0.0.1 frontend origins."
))
add_table(
    doc,
    ["Method", "Endpoint", "Purpose"],
    [
        ("GET", "/", "API status, documentation link, and route index"),
        ("GET", "/user", "Current workspace user"),
        ("GET", "/catalogs", "Available Unity Catalog catalogs"),
        ("GET", "/catalogs/{catalog}/schemas", "Schemas inside a catalog"),
        ("GET", "/catalogs/{catalog}/schemas/{schema}/tables", "Tables inside a schema"),
        ("GET", "/catalogs/{catalog}/schemas/{schema}/tables/{table}/data", "Read-only row preview"),
    ],
    [950, 5200, 3210],
)

doc.add_heading("Implementation safeguards", level=2)
for item in (
    "Path components are URL-encoded by the frontend.",
    "SQL identifiers are enclosed in backticks, with embedded backticks escaped.",
    "The preview limit is validated by FastAPI between 1 and 50 rows.",
    "Databricks statement failures return a readable error instead of an unhandled exception.",
    "API responses include success indicators and structured column/row payloads.",
):
    add_bullet(doc, item)

doc.add_heading("4. Validation completed", level=1)
p = doc.add_paragraph()
set_font(p.add_run(
    "The implementation was checked at the build, syntax, HTTP, CORS, metadata, and live query layers."
))
add_table(
    doc,
    ["Validation", "Observed result"],
    [
        ("Frontend production build", "Vite build completed successfully"),
        ("Frontend lint", "ESLint completed with no reported errors"),
        ("Backend syntax", "Python compilation passed for app.py and databricks_client.py"),
        ("Frontend availability", "HTTP 200 at http://127.0.0.1:5173"),
        ("Backend availability", "HTTP 200 at http://127.0.0.1:8000"),
        ("CORS", "Frontend origin accepted by backend"),
        ("Metadata drill-down", "system → access returned 8 tables"),
        ("Live data query", "system.access.assistant_events succeeded with 7 columns"),
    ],
    [3000, 6360],
)

doc.add_heading("5. How to run and use the application", level=1)
doc.add_heading("Start the backend", level=2)
add_step(doc, "Open a PowerShell terminal.", "Change directory to the backend folder.")
add_step(doc, "Start FastAPI.", r"Run .\venv\Scripts\uvicorn.exe app:app --reload")
add_step(doc, "Confirm API health.", "Open http://127.0.0.1:8000 or the interactive docs at /docs.")

doc.add_heading("Start the frontend", level=2)
add_step(doc, "Open a second PowerShell terminal.", "Change directory to the frontend folder.")
add_step(doc, "Start Vite.", "Run npm run dev -- --host 127.0.0.1")
add_step(doc, "Open the application.", "Navigate to http://127.0.0.1:5173.")

doc.add_heading("Explore workspace data", level=2)
add_step(doc, "Sync catalogs.", "Select the Sync catalogs action on the overview page.")
add_step(doc, "Expand a catalog.", "Its schemas appear directly beneath the catalog row.")
add_step(doc, "Expand a schema.", "Its tables and views appear beneath the schema row.")
add_step(doc, "Preview a table.", "Select a table to run a read-only query and show the first 25 rows.")

doc.add_heading("6. Files changed", level=1)
add_table(
    doc,
    ["File", "Responsibility"],
    [
        ("frontend/src/App.jsx", "Application shell, API state, nested explorer, data grid"),
        ("frontend/src/App.css", "Professional layout, responsive behavior, all interaction states"),
        ("frontend/src/index.css", "Global typography and page foundation"),
        ("backend/app.py", "FastAPI routes, CORS, root status response, preview limit validation"),
        ("backend/databricks_client.py", "Catalog/schema/table listing and SQL data preview"),
    ],
    [3150, 6210],
)

doc.add_heading("7. Current configuration and operational notes", level=1)
for item in (
    "The backend reads DATABRICKS_HOST and DATABRICKS_TOKEN from backend/.env.",
    "Data previews require access to at least one Databricks SQL warehouse and SELECT permission on the chosen table.",
    "The workspace identity must also have Unity Catalog metadata visibility for catalogs, schemas, and tables.",
    "The root backend URL is an API status page; the visual application always runs on the frontend URL.",
    "This is a local development configuration. Production deployment should move API URLs to environment variables and use managed authentication.",
):
    add_bullet(doc, item)

doc.add_heading("Final state", level=1)
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(4)
p.paragraph_format.space_after = Pt(0)
set_font(p.add_run(
    "The Privacy Control Center is running as an integrated local application with a polished Databricks-inspired interface and verified end-to-end navigation from Unity Catalog metadata to live table rows."
), size=11.5, color=GREEN, bold=True)

doc.core_properties.title = "Privacy & Compliance Control Center - Development Summary"
doc.core_properties.subject = "Summary of frontend and backend work completed June 30, 2026"
doc.core_properties.author = "Privacy Control Center Project"
doc.core_properties.keywords = "Databricks, Unity Catalog, FastAPI, React, Privacy Control Center"

OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUT)
print(OUT)
