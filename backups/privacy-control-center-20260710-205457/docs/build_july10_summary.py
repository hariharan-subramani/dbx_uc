from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "Privacy_Control_Center_July_10_2026_Work_Summary.docx"

INK = RGBColor(17, 24, 39)
MUTED = RGBColor(75, 85, 99)
RED = RGBColor(185, 28, 28)


def set_run(run, size=10.5, bold=False, color=INK):
    run.font.name = "Aptos"
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color


def para(doc, text="", size=10.5, bold=False, color=INK, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.08
    r = p.add_run(text)
    set_run(r, size=size, bold=bold, color=color)
    return p


def heading(doc, text, level=1):
    p = doc.add_heading(level=level)
    p.paragraph_format.space_before = Pt(10 if level == 1 else 6)
    p.paragraph_format.space_after = Pt(5)
    if not p.runs:
        p.add_run(text)
    p.runs[0].text = text
    for run in p.runs:
        set_run(run, size=16 if level == 1 else 12.5, bold=True, color=RED if level == 1 else INK)
    return p


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    set_run(r, size=10.5, color=INK)
    return p


def table(doc, rows):
    tbl = doc.add_table(rows=1, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = "Table Grid"
    hdr = tbl.rows[0].cells
    hdr[0].text = "Area"
    hdr[1].text = "Summary"
    for cell in hdr:
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for p in cell.paragraphs:
            for r in p.runs:
                set_run(r, bold=True, color=RED)
    for left, right in rows:
        cells = tbl.add_row().cells
        cells[0].text = left
        cells[1].text = right
        for cell in cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            for p in cell.paragraphs:
                for r in p.runs:
                    set_run(r, size=10)
    return tbl


doc = Document()
section = doc.sections[0]
section.start_type = WD_SECTION.NEW_PAGE
section.top_margin = Inches(0.65)
section.bottom_margin = Inches(0.65)
section.left_margin = Inches(0.75)
section.right_margin = Inches(0.75)

styles = doc.styles
styles["Normal"].font.name = "Aptos"
styles["Normal"].font.size = Pt(10.5)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Privacy & Compliance Control Center")
set_run(r, size=22, bold=True, color=RED)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("July 10, 2026 Work Summary")
set_run(r, size=13, bold=True, color=INK)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run(
    "Governance management, Unity Catalog volumes, group exploration, permission workflows, validation, and backup."
)
set_run(r, size=10.5, color=MUTED)

heading(doc, "Executive Summary")
para(
    doc,
    "Today the Privacy & Compliance Control Center was expanded from a governance dashboard into a stronger governance management platform. The existing React, FastAPI, and Databricks REST API architecture was preserved, and the current UI style, navigation, spacing, and color palette were kept intact.",
)
para(
    doc,
    "The main outcome is that administrators can now explore more Unity Catalog assets, inspect users and groups, and manage permissions through safer Grant, Edit, and Remove workflows. The application also now validates whether a principal can receive Unity Catalog permissions before attempting a grant, preventing unsupported workspace-only groups from producing confusing Databricks errors.",
)

heading(doc, "Major Enhancements")
table(
    doc,
    [
        ("Navigation cleanup", "Removed the old standalone Group Management module and kept the top navigation focused on Data Explorer and User Access Explorer."),
        ("User Access Explorer", "Added Group Explorer as a third tab beside Search User and Compare Users, with search-driven loading for Databricks groups."),
        ("Search behavior", "Updated user and group search flows so existing users or groups are not displayed before the administrator starts searching."),
        ("Unity Catalog volumes", "Added Volume support under Workspace, Catalog, Schema, and Table navigation. Volumes now load by schema and display details, permissions, exports, and binding information where available."),
        ("Permission management", "Made Grant Access, Edit Permission, and Remove Permission functional for catalogs, schemas, tables, and volumes."),
        ("Audit readiness", "Added a backend audit service and audit endpoint to record permission operations for later UI display."),
    ],
)

heading(doc, "Backend Work")
bullet(doc, "Added Unity Catalog volume REST integration for listing volumes, loading metadata, reading permissions, and checking binding information where Databricks exposes it.")
bullet(doc, "Added unified permission management endpoints: POST /permissions/grant, PATCH /permissions/edit, DELETE /permissions/remove, and GET /permissions/audit.")
bullet(doc, "Added POST /permissions/validate-principal to verify principal grantability before permission writes are attempted.")
bullet(doc, "Implemented principal classification for users, groups, workspace-only groups, account-level groups where available, and service-principal readiness.")
bullet(doc, "Added duplicate grant prevention so an already assigned privilege is not sent again unnecessarily.")
bullet(doc, "Improved group resolution so a single partial group match can be resolved while ambiguous matches produce a clear message.")

heading(doc, "Frontend Work")
bullet(doc, "Added the Group Explorer tab inside User Access Explorer without redesigning the surrounding interface.")
bullet(doc, "Added a Volume selector beneath the existing Table selector and wired it into the right-side object details panel.")
bullet(doc, "Reused the existing object details card, permission table, search controls, buttons, accordions, and export patterns for volumes and permission management.")
bullet(doc, "Updated the Grant Access modal with principal type selection, Databricks-backed autocomplete, valid privilege selection, duplicate privilege blocking, and automatic permission refresh.")
bullet(doc, "Added grantability status messaging in the modal: Grantable, Workspace Group, or Principal Not Found.")
bullet(doc, "Added principal dropdown badges so administrators can see whether a user or group is likely grantable before selection.")

heading(doc, "Grantability Validation")
para(
    doc,
    "A key safety improvement was added to stop unsupported Unity Catalog grants before they happen. The app now validates the selected principal against Databricks REST API data before enabling Grant Access.",
)
bullet(doc, "Grantable users can proceed normally.")
bullet(doc, "Workspace-only groups are blocked with a clear explanation that they cannot receive Unity Catalog permissions.")
bullet(doc, "Missing principals show a Principal Not Found state.")
bullet(doc, "The backend returns structured validation data so the frontend can display meaningful status instead of raw REST API errors.")

heading(doc, "Validation and Runtime Checks")
bullet(doc, "Ran backend Python compile checks after backend changes.")
bullet(doc, "Ran frontend production builds after React and CSS changes.")
bullet(doc, "Restarted the local FastAPI server after adding new backend routes.")
bullet(doc, "Verified the new validate-principal endpoint with the admins group and confirmed it returns Workspace Group with grantable set to false.")
bullet(doc, "Verified the volume endpoint after restart so hello.test.vol1 could load from Databricks.")

heading(doc, "Backup")
para(doc, "A backup of today's completed work was saved in the existing project backup location.")
bullet(doc, r"C:\Users\harih\Privacy-Control-Center\backups\privacy-control-center-20260710-112739")
bullet(doc, r"C:\Users\harih\Privacy-Control-Center\backups\privacy-control-center-20260710-112739.zip")

heading(doc, "Current Result")
para(
    doc,
    "The application now supports Data Explorer, User Access Explorer, Group Explorer, Volume Explorer, and object-level permission management across catalogs, schemas, tables, and volumes. It remains visually consistent with the existing application while adding safer governance operations and clearer Databricks principal handling.",
)

footer = section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = footer.add_run("Privacy & Compliance Control Center - July 10, 2026")
set_run(r, size=8.5, color=MUTED)

doc.core_properties.title = "Privacy & Compliance Control Center - July 10, 2026 Work Summary"
doc.core_properties.subject = "Summary of governance management enhancements completed on July 10, 2026"
doc.core_properties.author = "Privacy Control Center Project"
doc.core_properties.keywords = "Databricks, Unity Catalog, FastAPI, React, Governance, Permissions"

doc.save(OUT)
print(OUT)
