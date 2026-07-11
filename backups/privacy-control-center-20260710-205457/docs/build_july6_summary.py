from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path(__file__).with_name("Privacy_Compliance_Control_Center_July_6_2026_Summary.docx")

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "202A31"
MUTED = "66727B"
LIGHT = "F2F4F7"
ACCENT = "D53F3F"
GREEN = "2B7A57"


def set_font(run, size=11, color=INK, bold=False, italic=False):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths[index]
            cell.width = Inches(width / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_page_field(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])


def add_bullet(doc, text):
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.left_indent = Inches(0.5)
    paragraph.paragraph_format.first_line_indent = Inches(-0.25)
    paragraph.paragraph_format.space_after = Pt(5)
    set_font(paragraph.add_run(text))


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)

    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        set_cell_shading(cell, LIGHT)
        paragraph = cell.paragraphs[0]
        paragraph.paragraph_format.space_after = Pt(0)
        set_font(paragraph.add_run(header), size=10, color=DARK_BLUE, bold=True)

    for row_data in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row_data):
            paragraph = cells[index].paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.05
            set_font(paragraph.add_run(str(value)), size=9.5)

    set_table_geometry(table, widths)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(0)
    return table


def add_callout(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_geometry(table, [9360])
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, "F8FAFC")
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(2)
    set_font(paragraph.add_run(title), size=10.5, color=DARK_BLUE, bold=True)
    paragraph.add_run("\n")
    set_font(paragraph.add_run(body), size=10.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.header_distance = Inches(0.492)
section.footer_distance = Inches(0.492)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
normal.font.size = Pt(11)
normal.font.color.rgb = RGBColor.from_string(INK)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.1

for name, size, color, before, after in (
    ("Heading 1", 16, BLUE, 16, 8),
    ("Heading 2", 13, BLUE, 12, 6),
    ("Heading 3", 12, DARK_BLUE, 8, 4),
):
    style = styles[name]
    style.font.name = "Calibri"
    style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = RGBColor.from_string(color)
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.keep_with_next = True

header = section.header.paragraphs[0]
header.alignment = WD_ALIGN_PARAGRAPH.LEFT
set_font(header.add_run("PRIVACY & COMPLIANCE CONTROL CENTER  /  DAILY SUMMARY"), size=8.5, color=MUTED, bold=True)

footer = section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
set_font(footer.add_run("6 Jul 2026  |  Page "), size=8.5, color=MUTED)
add_page_field(footer)

label = doc.add_paragraph()
label.paragraph_format.space_after = Pt(3)
set_font(label.add_run("PROJECT UPDATE"), size=9, color=ACCENT, bold=True)

title = doc.add_paragraph()
title.paragraph_format.space_after = Pt(4)
set_font(title.add_run("Privacy & Compliance Control Center"), size=25, color=INK, bold=True)

subtitle = doc.add_paragraph()
subtitle.paragraph_format.space_after = Pt(14)
set_font(
    subtitle.add_run("Summary of today's work: Databricks REST API migration, dynamic metadata, permissions, catalog binding, and Excel exports"),
    size=12.5,
    color=MUTED,
)

metadata = [
    ("Project", "Privacy-Control-Center"),
    ("Date", "6 Jul 2026"),
    ("Frontend", "React 19 + Vite"),
    ("Backend", "FastAPI + Databricks REST APIs"),
    ("Current app URL", "http://127.0.0.1:5174/"),
    ("Current API URL", "http://127.0.0.1:8001"),
]
for key, value in metadata:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(2)
    set_font(paragraph.add_run(f"{key}: "), size=10.5, bold=True)
    set_font(paragraph.add_run(value), size=10.5)

add_callout(
    doc,
    "Key architecture change",
    "The application moved away from relying on the Databricks SDK-style service layer for active functionality and now consumes Databricks REST API responses through FastAPI. This makes the frontend ready for real Unity Catalog metadata, permissions, workspace binding, and table responses instead of mock or hardcoded data.",
)

doc.add_heading("Executive Summary", level=1)
paragraph = doc.add_paragraph()
set_font(paragraph.add_run(
    "Today we converted the Privacy & Compliance Control Center into a more dynamic governance application. The UI layout was preserved, but the data flow was changed so selected catalogs, schemas, and tables refresh from backend APIs. We also fixed real catalog permission parsing, renamed Catalog Bundles to Catalog Binding, added friendly workspace naming, and introduced enterprise-style Excel exports."
))

doc.add_heading("Major Work Completed", level=1)
for item in (
    "Replaced frontend hardcoded metadata with live backend responses for catalogs, schemas, and tables.",
    "Changed the application direction from SDK-centered implementation to Databricks REST API consumption through FastAPI.",
    "Added dynamic catalog metadata, schema metadata, table metadata, permissions, catalog binding, schema objects, and table statistics flows.",
    "Fixed catalog permissions by parsing Databricks REST API privilege_assignments responses.",
    "Added loading states and friendly empty/error messages for unavailable metadata, permissions, and binding data.",
    "Kept the existing UI layout and styling while improving functionality.",
):
    add_bullet(doc, item)

doc.add_heading("SDK to REST API Migration", level=1)
paragraph = doc.add_paragraph()
set_font(paragraph.add_run(
    "The backend now acts as a REST API adapter between the React frontend and Databricks. The FastAPI service calls Databricks REST endpoints such as Unity Catalog catalogs, schemas, tables, permissions, workspace bindings, and SQL statements. This creates a clearer contract for the frontend and avoids frontend mock data."
))

add_table(
    doc,
    ["Area", "Before", "Now"],
    [
        ("Catalog list", "SDK/service abstraction and static UI assumptions", "REST-backed /catalogs response"),
        ("Metadata", "Partially hardcoded in frontend", "Fetched from FastAPI endpoints per selected object"),
        ("Permissions", "Empty because parser missed Databricks response shape", "Parsed from REST privilege_assignments"),
        ("Catalog workspace access", "Mock bundle values", "REST-backed catalog binding response"),
        ("Exports", "Not available", "SheetJS/FileSaver Excel workbooks from currently loaded UI data"),
    ],
    [1800, 3780, 3780],
)

doc.add_heading("Frontend Improvements", level=1)
add_table(
    doc,
    ["Feature", "What Changed"],
    [
        ("Workspace selector", "Shows a friendly workspace label while keeping host/ID internally."),
        ("Catalog selection", "Automatically reloads metadata, permissions, schemas, and catalog binding."),
        ("Schema selection", "Automatically reloads schema metadata, permissions, tables, and schema objects."),
        ("Table selection", "Automatically reloads table metadata, permissions, and statistics."),
        ("Permissions table", "Displays Principal, Type, and Privileges with badge values as plain text for export."),
        ("Search", "User search filters only users; group search filters only groups; both are case-insensitive partial matches."),
    ],
    [2300, 7060],
)

doc.add_heading("Backend Improvements", level=1)
for item in (
    "Added /workspace endpoint for friendly workspace information.",
    "Kept existing FastAPI endpoint structure and added only what was needed.",
    "Updated Unity Catalog permissions calls to use singular securable types: catalog, schema, and table.",
    "Normalized Databricks principal and privilege responses for frontend display.",
    "Handled Databricks Free Edition or unsupported endpoint cases with clear messages.",
    "Derived table statistics from real table metadata where available.",
):
    add_bullet(doc, item)

doc.add_heading("Permissions Fix", level=1)
paragraph = doc.add_paragraph()
set_font(paragraph.add_run(
    "The permissions section originally showed 'No permissions found' even when Databricks returned real grants. The root cause was that Databricks returned permissions under privilege_assignments, while the backend parser only checked permission_assignments. The parser now supports both shapes."
))

add_table(
    doc,
    ["Databricks Response", "Normalized UI Output"],
    [
        ("principal: account users", "Principal: account users"),
        ("privileges: BROWSE", "Privileges: BROWSE"),
        ("principal_id returned by API", "Not shown in current UI, available for future use"),
        ("No assignments returned", "No permissions found"),
    ],
    [4680, 4680],
)

doc.add_heading("Excel Export Feature", level=1)
paragraph = doc.add_paragraph()
set_font(paragraph.add_run(
    "We added enterprise-grade Excel exports using SheetJS and FileSaver.js. The implementation is reusable so future schema, table, policy, audit log, and compliance report exports can use the same helper pattern."
))
add_table(
    doc,
    ["Export", "Workbook", "Sheets"],
    [
        ("Catalog Details card", "Catalog_<CatalogName>.xlsx", "Catalog Details, Permissions, Catalog Binding"),
        ("Permissions accordion", "CatalogPermissions_<CatalogName>.xlsx", "Permissions"),
        ("Catalog Binding accordion", "CatalogBinding_<CatalogName>.xlsx", "Catalog Binding"),
    ],
    [2600, 3180, 3580],
)

doc.add_heading("Files Changed", level=1)
add_table(
    doc,
    ["File", "Purpose"],
    [
        ("backend/app.py", "FastAPI routes and REST API endpoint wiring"),
        ("backend/services/databricks_api.py", "Databricks REST API requests, normalization, metadata, permissions, binding, statistics"),
        ("frontend/src/components/Explorer.jsx", "Dynamic selection and API refresh flow"),
        ("frontend/src/components/ObjectDetailsCard.jsx", "Details display, permissions, catalog binding, and export buttons"),
        ("frontend/src/components/PermissionTable.jsx", "Principal, Type, Privileges display and search behavior"),
        ("frontend/src/components/WorkspaceSelector.jsx", "Friendly workspace label display"),
        ("frontend/src/utils/excelExport.js", "Reusable SheetJS/FileSaver export helpers"),
        ("frontend/src/App.css", "Small export button styling without changing layout"),
    ],
    [3300, 6060],
)

doc.add_heading("Verification Completed", level=1)
for item in (
    "Frontend lint passed with npm run lint.",
    "Frontend production build passed with npm run build.",
    "Backend syntax checks passed for app.py and databricks_api.py.",
    "Verified /workspace returns a friendly Databricks Workspace label on the updated backend.",
    "Verified /catalogs/fortest/permissions returns account users as Group with BROWSE privilege.",
    "Verified /catalogs/fortest/bundles returns real catalog binding data.",
    "Confirmed the active working app is http://127.0.0.1:5174/ connected to backend port 8001.",
):
    add_bullet(doc, item)

doc.add_heading("How to Restart Later", level=1)
add_table(
    doc,
    ["Terminal", "Command"],
    [
        ("Backend", r"cd C:\Users\harih\Privacy-Control-Center\backend"),
        ("Backend", r".\venv\Scripts\python.exe -m uvicorn app:app --host 127.0.0.1 --port 8001"),
        ("Frontend", r"cd C:\Users\harih\Privacy-Control-Center\frontend"),
        ("Frontend", r"$env:VITE_API_BASE='http://127.0.0.1:8001'"),
        ("Frontend", "npm run dev -- --host 127.0.0.1 --port 5174"),
        ("Browser", "Open http://127.0.0.1:5174/"),
    ],
    [1800, 7560],
)

doc.add_heading("Current Status", level=1)
paragraph = doc.add_paragraph()
set_font(paragraph.add_run(
    "The application now behaves more like a Databricks Catalog Explorer for governance review: selecting a catalog dynamically loads metadata, permissions, and catalog binding; users can search permissions and export audit-ready Excel workbooks."
), size=11.5, color=GREEN, bold=True)

doc.core_properties.title = "Privacy & Compliance Control Center - 6 Jul 2026 Summary"
doc.core_properties.subject = "Databricks REST API migration and export feature summary"
doc.core_properties.author = "Privacy Control Center Project"
doc.core_properties.keywords = "Databricks REST API, Unity Catalog, FastAPI, React, Excel export, governance"

OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUT)
print(OUT)
