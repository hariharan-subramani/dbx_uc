from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "Privacy_Control_Center_July_13_2026_Work_Summary.docx"
BACKUP = "backups/privacy-control-center-oauth-live-20260713-100148.zip"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(31, 41, 55)
MUTED = RGBColor(91, 103, 117)
LIGHT = "F2F4F7"
PALE_BLUE = "E8EEF5"
WHITE = "FFFFFF"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[idx]))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(widths_dxa[idx] / 1440)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)


def add_table(doc, headers, rows, widths_dxa):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = header
        set_cell_shading(cell, PALE_BLUE)
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = DARK_BLUE
    tr_pr = table.rows[0]._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)
    for values in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(values):
            cells[idx].text = str(value)
            if len(table.rows) % 2 == 1:
                set_cell_shading(cells[idx], LIGHT)
    set_table_geometry(table, widths_dxa)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.167
    p.add_run(text)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.167
    p.add_run(text)
    return p


def add_callout(doc, label, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.08)
    p.paragraph_format.right_indent = Inches(0.08)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "F4F6F9")
    p_pr.append(shd)
    borders = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:color"), "2E74B5")
    left.set(qn("w:space"), "6")
    borders.append(left)
    p_pr.append(borders)
    lead = p.add_run(f"{label}: ")
    lead.bold = True
    lead.font.color.rgb = DARK_BLUE
    p.add_run(text)


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.header_distance = Inches(0.492)
section.footer_distance = Inches(0.492)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)
normal.font.color.rgb = INK
normal.paragraph_format.space_before = Pt(0)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.10

for name, size, color, before, after in (
    ("Heading 1", 16, BLUE, 16, 8),
    ("Heading 2", 13, BLUE, 12, 6),
    ("Heading 3", 12, DARK_BLUE, 8, 4),
):
    style = styles[name]
    style.font.name = "Calibri"
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = color
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.keep_with_next = True

header = section.header.paragraphs[0]
header.text = "Privacy & Compliance Control Center | Technical Change Summary"
header.alignment = WD_ALIGN_PARAGRAPH.LEFT
header.runs[0].font.size = Pt(9)
header.runs[0].font.color.rgb = MUTED

footer = section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run = footer.add_run("Internal technical record | 13 July 2026")
run.font.size = Pt(9)
run.font.color.rgb = MUTED

# Memo masthead
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(16)
p.paragraph_format.space_after = Pt(4)
r = p.add_run("DAILY WORK SUMMARY")
r.bold = True
r.font.size = Pt(23)
r.font.color.rgb = INK

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(16)
r = p.add_run("Privacy & Compliance Control Center - 13 July 2026")
r.font.size = Pt(14)
r.font.color.rgb = MUTED

metadata = [
    ("Application", "Privacy & Compliance Control Center"),
    ("Environment", "Single Databricks workspace"),
    ("Authentication", "OAuth 2.0 client credentials (machine-to-machine)"),
    ("Completed", "13 July 2026"),
    ("Backup", BACKUP),
]
for label, value in metadata:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    a = p.add_run(f"{label}: ")
    a.bold = True
    p.add_run(value)

doc.add_paragraph()
add_callout(
    doc,
    "Outcome",
    "The application now authenticates successfully to the real Databricks workspace, refreshes OAuth tokens automatically, serves live governance data, and returns verified catalog, schema, and table access in User Access Explorer.",
)

doc.add_heading("1. Executive Summary", level=1)
doc.add_paragraph(
    "The control center was moved from legacy or demo-style workspace connectivity to a live, single-workspace Databricks integration. The final design intentionally does not include multi-workspace discovery. Backend credentials remain in the server-side environment file and are exchanged for short-lived OAuth access tokens. The frontend receives workspace metadata and governance results only; it never receives the client secret or access token."
)
doc.add_paragraph(
    "During verification, OAuth initially failed because account-level variable names and an invalid credential pair were used for a workspace-level flow. After the environment variables were corrected and a valid workspace service-principal credential pair was supplied, the token exchange and workspace API verification succeeded."
)

doc.add_heading("2. Final Credential Configuration", level=1)
add_table(
    doc,
    ["Environment variable", "Purpose"],
    [
        ("DATABRICKS_HOST", "HTTPS URL of the one connected Databricks workspace."),
        ("DATABRICKS_WORKSPACE_NAME", "Optional friendly name displayed in the disabled workspace selector."),
        ("DATABRICKS_CLIENT_ID", "Application/client ID of the workspace-assigned Databricks service principal."),
        ("DATABRICKS_CLIENT_SECRET", "OAuth secret used only by the backend token exchange."),
    ],
    [2700, 6660],
)
add_callout(
    doc,
    "Security",
    "No secret value, bearer token, or full credential pair is reproduced in this document. Credentials must remain in backend/.env, must not be committed to source control, and should be rotated if exposed.",
)

doc.add_heading("3. OAuth Connection Flow", level=1)
for text in (
    "At backend startup, validate that the workspace host, client ID, and client secret are configured.",
    "POST a client-credentials request to the workspace /oidc/v1/token endpoint with scope all-apis.",
    "Cache the returned access token in backend memory and refresh it before expiration.",
    "Attach the bearer token automatically to every Databricks REST API request.",
    "If Databricks returns HTTP 401, refresh the token once and retry the original request.",
    "Verify the connection with the live Workspace API before presenting the workspace as connected.",
):
    add_number(doc, text)

doc.add_heading("4. Backend Changes", level=1)
for text in (
    "Introduced services/workspace_auth.py as the single source for configuration validation, token issuance, refresh, and workspace identity.",
    "Removed the Databricks Accounts API and workspace-selection endpoints; GET /workspace is the only workspace route.",
    "Centralized OAuth bearer-token injection in the existing Databricks REST request helper, preserving catalogs, schemas, tables, volumes, metadata, permissions, bindings, identity explorers, comparison, permission mutations, and exports.",
    "Removed the hardcoded fallback user so failed identity calls cannot appear as valid demo data.",
    "Set governance response-cache duration to zero so requests retrieve current workspace information.",
    "Updated test_connection.py to verify OAuth and the workspace REST API without printing credentials or tokens.",
    "Normalized network and OAuth failures into a controlled workspace-authentication error so startup diagnostics remain clear.",
):
    add_bullet(doc, text)

doc.add_heading("5. Frontend Changes", level=1)
for text in (
    "The Data Explorer loads the verified workspace from GET /workspace and then loads catalogs.",
    "The workspace control remains visually consistent with the existing design but is disabled because only one workspace is supported.",
    "Connection failure is displayed as: Unable to connect to Databricks Workspace.",
    "The client secret and OAuth access token are never sent to or stored by the frontend.",
):
    add_bullet(doc, text)

doc.add_heading("6. User Access Explorer Issue and Repair", level=1)
doc.add_paragraph(
    "After the workspace connected, user search returned real SCIM users but catalog access displayed zero. Live diagnostics confirmed that users and catalogs were available, while the access scan stopped after seven objects and returned partial results."
)
add_table(
    doc,
    ["Root cause", "Correction"],
    [
        ("Depth-first catalog traversal", "Scan permissions for every catalog before descending into schemas, tables, and volumes."),
        ("15-second scan budget", "Increase the default budget to 60 seconds, request timeout to 10 seconds, and object limit to 1,000."),
        ("Missing built-in membership", "Treat every assigned workspace identity as a member of the Databricks built-in account users principal."),
        ("Duplicate effective entries", "Merge direct and inherited grants by securable name while combining privileges, sources, and principals."),
    ],
    [3200, 6160],
)

doc.add_heading("7. Live Verification Results", level=1)
add_table(
    doc,
    ["Verification", "Result"],
    [
        ("OAuth client-credentials exchange", "Passed"),
        ("Workspace REST API verification", "Passed"),
        ("Connected deployment", "dbc-d5d034c9-7400"),
        ("Live SCIM users returned", "4"),
        ("Live catalogs returned", "8"),
        ("Known-user effective catalogs after repair", "5"),
        ("Known-user schemas returned during scan", "2"),
        ("Known-user tables returned during scan", "28"),
        ("Objects scanned in verification", "39"),
        ("Python compilation", "Passed"),
        ("Vite production build", "Passed"),
        ("Frontend and backend HTTP checks", "Passed on ports 5173 and 8000"),
    ],
    [4200, 5160],
)
doc.add_paragraph(
    "Large workspaces can still reach the configured scan limit for lower-level objects. Catalog permissions are now scanned first, so catalog access is not incorrectly reported as zero even when the deeper scan is partial."
)

doc.add_heading("8. Operating Procedure", level=1)
for text in (
    "Maintain only the four DATABRICKS_* variables described in Section 2.",
    "Restart the backend after changing .env because configuration is loaded at process startup.",
    "Run python test_connection.py from the backend directory to validate OAuth safely.",
    "Start FastAPI on port 8000 and Vite on port 5173, then confirm GET /workspace returns connected=true.",
    "Use User Access Explorer to search a real SCIM user and review catalogs first; treat a partial-scan warning as incomplete lower-level coverage, not as zero catalog access.",
):
    add_number(doc, text)

doc.add_heading("9. Backup Record", level=1)
doc.add_paragraph(
    "A timestamped archive was created after the successful OAuth connection and User Access Explorer repair. Generated dependency and build caches were excluded. The archive contains the saved backend, frontend, documentation, and environment configuration state."
)
add_callout(doc, "Backup file", BACKUP)

doc.add_heading("10. Rollbacks and Final Service State", level=1)
doc.add_paragraph(
    "A bounded concurrent permission-scan experiment and an additional incomplete-scan frontend alert were introduced briefly, then reverted at the user's request. The earlier catalog-first scan, built-in account users matching, merged effective permissions, OAuth integration, and single-workspace design were preserved."
)
doc.add_paragraph(
    "The backend was later found stopped, which caused the frontend to show both Workspace connection could not be verified and Unable to connect to Databricks Workspace. The backend was restarted and the issue was confirmed to be process availability rather than credential failure."
)
add_table(
    doc,
    ["Final check", "Status"],
    [
        ("Frontend on port 5173", "Running and returning HTTP 200"),
        ("Backend on port 8000", "Running and returning HTTP 200"),
        ("Databricks OAuth", "Connected"),
        ("Workspace API", "connected=true"),
        ("SCIM user API", "Successful; 4 live users returned"),
    ],
    [4200, 5160],
)

doc.add_heading("11. Recommended Follow-up", level=1)
for text in (
    "Set DATABRICKS_WORKSPACE_NAME to a friendly operational name if the deployment identifier is not suitable for the UI.",
    "Rotate any client secret that was previously shared outside the secure environment file.",
    "Add .env to source-control ignore rules and use a managed secret store for production deployment.",
    "Consider a parallelized permission scan or an asynchronous scan job if the workspace grows beyond the current 60-second interactive budget.",
    "Add automated integration tests for OAuth refresh, catalog-first traversal, account users inheritance, and merged effective permissions.",
):
    add_bullet(doc, text)

doc.core_properties.title = "Privacy Control Center - 13 July 2026 Work Summary"
doc.core_properties.subject = "Databricks OAuth integration, live API verification, and User Access Explorer repair"
doc.core_properties.author = "Privacy & Compliance Control Center Engineering"
doc.core_properties.keywords = "Databricks, OAuth, FastAPI, React, Unity Catalog, governance"

doc.save(OUT)
print(OUT)
